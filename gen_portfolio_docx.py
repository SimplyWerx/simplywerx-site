#!/usr/bin/env python3
"""Generate a portfolio .docx from the _portfolio/ markdown files."""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).parent
PORTFOLIO_DIR = BASE / "_portfolio"
IMG_LARGE = BASE / "assets/img/optimized-large"
IMG_SMALL = BASE / "assets/img/optimized-small"
IMG_ORIG  = BASE / "assets/img/originals"
OUT = BASE / "SimplyWerx_Portfolio.docx"

def find_image(src: str) -> Path | None:
    """Resolve an image src attribute to a local file, trying common variants."""
    fname = Path(src).name
    stem = Path(fname).stem
    for folder in [IMG_LARGE, IMG_SMALL, IMG_ORIG, BASE]:
        for ext in ["", ".jpg", ".jpeg", ".png", ".webp", ".svg"]:
            candidate = folder / (fname if not ext else stem + ext)
            if candidate.exists():
                return candidate
    return None

def parse_front_matter(text: str):
    """Return (front_matter_dict_raw, body) from a Jekyll markdown file."""
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end == -1:
        return {}, text
    fm_raw = text[3:end].strip()
    body = text[end + 4:].strip()
    return fm_raw, body

def extract_fm_value(fm_raw: str, key: str) -> str:
    m = re.search(rf"^{key}:\s*[\"']?(.+?)[\"']?\s*$", fm_raw, re.MULTILINE)
    return m.group(1).strip() if m else ""

def extract_caption(fm_raw: str, field: str) -> str:
    m = re.search(rf"^\s+{field}:\s*[\"']?(.+?)[\"']?\s*$", fm_raw, re.MULTILINE)
    return m.group(1).strip() if m else ""

def add_image_to_doc(doc: Document, src: str, alt: str = ""):
    img_path = find_image(src)
    if img_path and img_path.suffix.lower() not in (".svg",):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(5.5))
            if alt:
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.italic = True
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        except Exception as e:
            doc.add_paragraph(f"[Image: {alt or src}]")
    else:
        doc.add_paragraph(f"[Image: {alt or src}]")

def parse_body(doc: Document, body: str):
    """Walk through the body HTML/markdown and add content to doc."""
    img_re   = re.compile(r'<img\s[^>]*src=["\']([^"\']+)["\'][^>]*(?:alt=["\']([^"\']*)["\'])?[^>]*/>', re.DOTALL)
    video_re = re.compile(r'<(?:iframe|video)[^>]*(?:src|data-src)=["\']([^"\']+)["\'][^>]*>', re.DOTALL)
    list_re  = re.compile(r'\{:\s*\.list-inline\s*\}')

    # Split on HTML tags we care about
    tokens = re.split(r'(<img\b[^>]*/?>|<iframe\b[^>]*>.*?</iframe>|<video\b.*?</video>|\{:.*?\})', body, flags=re.DOTALL)

    for token in tokens:
        token = token.strip()
        if not token:
            continue

        # Image
        m = img_re.match(token)
        if m:
            add_image_to_doc(doc, m.group(1), m.group(2) or "")
            continue

        # Video / iframe
        if re.match(r'<(?:iframe|video)', token, re.IGNORECASE):
            vm = re.search(r'(?:src|data-src)=["\']([^"\']+)["\']', token)
            label = vm.group(1) if vm else "video"
            p = doc.add_paragraph(f"[Video: {label}]")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            continue

        # Liquid / Jekyll list-inline marker – skip
        if list_re.match(token):
            continue

        # Blockquote
        if token.startswith(">"):
            lines = [l.lstrip("> ").strip() for l in token.splitlines() if l.strip().lstrip("> ")]
            p = doc.add_paragraph("\n".join(lines), style="Quote")
            continue

        # Heading ##
        hm = re.match(r'^(#{1,3})\s+(.+)', token)
        if hm:
            level = len(hm.group(1))
            doc.add_heading(hm.group(2).strip(), level=level + 1)
            continue

        # Plain list items (- ...)
        if re.match(r'^-\s', token):
            for line in token.splitlines():
                lm = re.match(r'^-\s+(.+)', line.strip())
                if lm:
                    doc.add_paragraph(lm.group(1), style="List Bullet")
            continue

        # Stars in text – strip emoji-style stars
        token_clean = re.sub(r'⭐+', '', token).strip()
        # Strip remaining HTML tags
        token_clean = re.sub(r'<[^>]+>', '', token_clean).strip()
        if token_clean:
            doc.add_paragraph(token_clean)


def main():
    doc = Document()

    # Title page
    doc.add_heading("SimplyWerx Portfolio", 0)
    doc.add_paragraph("A selection of interactive exhibits and products developed by SimplyWerx Pte Ltd.")
    doc.add_page_break()

    files = sorted(PORTFOLIO_DIR.glob("*.md"))
    for md_file in files:
        text = md_file.read_text(encoding="utf-8")
        fm_raw, body = parse_front_matter(text)

        title    = extract_fm_value(fm_raw, "title")
        subtitle = extract_fm_value(fm_raw, "subtitle")
        image    = extract_fm_value(fm_raw, "image")
        alt      = extract_fm_value(fm_raw, "alt")
        client   = ""
        category = ""

        # Extract client/category from list-inline block
        for line in body.splitlines():
            cm = re.match(r'^-\s+Client:\s+(.+)', line.strip())
            if cm:
                client = cm.group(1)
            catm = re.match(r'^-\s+Category:\s+(.+)', line.strip())
            if catm:
                category = catm.group(1)

        doc.add_heading(title, level=1)
        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.runs[0].font.italic = True

        meta_parts = []
        if client:
            meta_parts.append(f"Client: {client}")
        if category:
            meta_parts.append(f"Category: {category}")
        if meta_parts:
            mp = doc.add_paragraph(" | ".join(meta_parts))
            mp.runs[0].font.bold = True
            mp.runs[0].font.size = Pt(10)

        doc.add_paragraph()  # spacer

        # Main hero image
        if image:
            add_image_to_doc(doc, image, alt)

        # Body content (strip list-inline lines, they're already parsed above)
        body_no_list = re.sub(r'\{:.*?\}', '', body, flags=re.DOTALL)
        body_no_list = re.sub(r'(?m)^-\s+(?:Client|Category):.+\n?', '', body_no_list)

        parse_body(doc, body_no_list)

        doc.add_page_break()

    doc.save(str(OUT))
    print(f"Saved: {OUT}")

if __name__ == "__main__":
    main()
